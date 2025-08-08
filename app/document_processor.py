import docx
import email
from email import policy
import logging
import os
from typing import List
import pdfplumber

from .document_downloader import DocumentDownloader
from .models import DocumentChunk
from .in_memory_vector_store import InMemoryVectorStore

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles the downloading, processing, and chunking of various document types."""

    def __init__(self, embedding_service, vector_store: InMemoryVectorStore):
        """Initializes the DocumentProcessor with an embedding service and a vector store."""
        self.embedding_service = embedding_service
        self.vector_store = vector_store

    async def process_document_from_url(self, url: str) -> List[DocumentChunk]:
        """Download, process document from URL and return chunks (no storage)"""
        downloader = DocumentDownloader()
        file_path = None

        try:
            logger.info(f"🔄 Processing document from URL: {url}")

            # Download document
            file_path, file_type = await downloader.download_document(url)

            # Extract text based on file type
            if file_type == 'pdf':
                text_content = self._extract_text_from_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                text_content = self._extract_text_from_docx(file_path)
            elif file_type == 'email':
                text_content = self._extract_text_from_email(file_path)
            else:
                logger.warning(f"⚠️ Unknown file type: {file_type}, trying PDF extraction")
                text_content = self._extract_text_from_pdf(file_path)

            if not text_content.strip():
                raise ValueError("No text content found in document")

            logger.info(f"📝 Extracted {len(text_content)} characters from {file_type} document")

            # Split into chunks
            chunks = self._split_text_into_chunks(text_content)
            logger.info(f"✂️ Split into {len(chunks)} chunks")

            # Generate embeddings for chunks
            embeddings = await self.embedding_service.generate_batch_embeddings(chunks)

            # Create DocumentChunk objects
            processed_chunks = []
            for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                chunk = DocumentChunk(
                    id=f"chunk_{i}_{hash(url)}",
                    text=chunk_text,
                    metadata={
                        "source": url,
                        "file_type": file_type,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "chunk_size": len(chunk_text)
                    },
                    embedding=embedding
                )
                processed_chunks.append(chunk)

            logger.info(f"✅ Processed {len(processed_chunks)} chunks from URL")
            return processed_chunks

        except Exception as e:
            logger.error(f"❌ Failed to process document from URL: {str(e)}")
            raise
        finally:
            # Cleanup downloaded file
            if file_path:
                downloader.cleanup_file(file_path)

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            logger.info(f"📄 Extracting text from DOCX: {file_path}")

            doc = docx.Document(file_path)
            text_content = ""

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"

            logger.info(f"📊 Extracted {len(text_content)} characters from DOCX")
            return text_content.strip()

        except Exception as e:
            logger.error(f"❌ Failed to extract text from DOCX: {str(e)}")
            raise

    def _extract_text_from_email(self, file_path: str) -> str:
        """Extract text from email file"""
        try:
            logger.info(f"📧 Extracting text from email: {file_path}")

            with open(file_path, 'rb') as f:
                msg = email.message_from_bytes(f.read(), policy=policy.default)

            text_content = ""

            # Extract headers
            text_content += f"From: {msg.get('From', 'Unknown')}\n"
            text_content += f"To: {msg.get('To', 'Unknown')}\n"
            text_content += f"Subject: {msg.get('Subject', 'No Subject')}\n"
            text_content += f"Date: {msg.get('Date', 'Unknown')}\n\n"

            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        text_content += part.get_content()
            else:
                if msg.get_content_type() == "text/plain":
                    text_content += msg.get_content()

            logger.info(f"📧 Extracted {len(text_content)} characters from email")
            return text_content.strip()

        except Exception as e:
            logger.error(f"❌ Failed to extract text from email: {str(e)}")
            raise

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Enhanced PDF extraction with better table handling"""
        try:
            text_content = ""

            with pdfplumber.open(file_path) as pdf:
                logger.info(f"📊 Processing {len(pdf.pages)} pages with table extraction")

                for page_num, page in enumerate(pdf.pages):
                    if page_num % 10 == 0:
                        logger.info(f"📄 Processing page {page_num + 1}/{len(pdf.pages)}")

                    # Extract regular text
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"

                    # Extract tables
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        logger.debug(f"📊 Found table {table_num + 1} on page {page_num + 1}")
                        text_content += f"\n[TABLE {table_num + 1} FROM PAGE {page_num + 1}]\n"

                        for row in table:
                            if row:  # Skip empty rows
                                # Clean and join row cells
                                clean_row = [str(cell).strip() if cell else "" for cell in row]
                                text_content += " | ".join(clean_row) + "\n"

                        text_content += "[END TABLE]\n\n"

            return text_content.strip()

        except Exception as e:
            logger.error(f"❌ Failed to extract text from PDF: {str(e)}")
            raise
    
    def _split_text_into_chunks(self, text_content: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        # This is a placeholder method for splitting text into chunks.
        # A more robust implementation would use a library like LangChain's text splitters.
        chunks = []
        if not text_content:
            return chunks

        text_length = len(text_content)
        start = 0
        while start < text_length:
            end = start + chunk_size
            chunk = text_content[start:end]
            chunks.append(chunk)
            start += chunk_size - overlap
            if start >= text_length:
                break
        return chunks